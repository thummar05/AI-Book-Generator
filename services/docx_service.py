import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# ==========================================
# SECTIONS THAT GET A PAGE BREAK BEFORE THEM
# ==========================================

PAGE_BREAK_SECTIONS = {
    "HALF_TITLE",
    "TITLE_PAGE",
    "COPYRIGHT",
    "DEDICATION",
    "EPIGRAPH",
    "TOC",
    "ACKNOWLEDGMENTS",
    "INTRODUCTION",
    "GLOSSARY",
    "ABOUT_AUTHOR",
    "BACK_COVER",
    "PREFACE",
    "AFTERWORD",
    "APPENDIX",
    "REFERENCES",
}

# ==========================================
# SECTION DISPLAY NAMES
# (None = no visible label, just page break)
# ==========================================

SECTION_LABELS = {
    "HALF_TITLE":      None,
    "TITLE_PAGE":      None,
    "COPYRIGHT":       "Copyright",
    "DEDICATION":      "Dedication",
    "EPIGRAPH":        None,
    "TOC":             "Table of Contents",
    "ACKNOWLEDGMENTS": "Acknowledgments",
    "INTRODUCTION":    "Introduction",
    "GLOSSARY":        "Glossary",
    "ABOUT_AUTHOR":    "About the Author",
    "BACK_COVER":      "Back Cover",
    "PREFACE":         "Preface",
    "AFTERWORD":       "Afterword",
    "APPENDIX":        "Appendix",
    "REFERENCES":      "References",
}


def _set_margins(document, margin_inches=1.0):
    """Set page margins on all sections."""
    for section in document.sections:
        section.top_margin    = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin   = Inches(margin_inches)
        section.right_margin  = Inches(margin_inches)


def _add_section_title(document, text):
    """Centered, large heading used for section labels."""
    p = document.add_heading(text, level=1)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(24)
    run.font.bold = True


def _add_book_title(document, text):
    """Big centered title for HALF_TITLE / TITLE_PAGE."""
    p = document.add_heading(text, level=1)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(28)
    run.font.bold = True


def _add_chapter_label(document, chapter_num):
    """Small grey 'Chapter N' label above the chapter title."""
    p = document.add_paragraph()
    run = p.add_run(f"Chapter {chapter_num}")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_after = Pt(2)


def _add_chapter_title(document, text):
    """Large left-aligned chapter title."""
    p = document.add_heading(text, level=2)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.runs[0]
    run.font.size = Pt(22)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(20)


def _add_section_heading(document, text):
    """Bold left-aligned heading inside a chapter."""
    p = document.add_heading(text, level=3)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.runs[0]
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)


def _add_body(document, text):
    """Standard body paragraph."""
    p = document.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)


def _add_toc_entry(document, text):
    """TOC line — normal style, no heading promotion."""
    p = document.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)


def _add_epigraph(document, text):
    """Italic centered epigraph text."""
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.italic = True
    p.paragraph_format.left_indent  = Inches(0.75)
    p.paragraph_format.right_indent = Inches(0.75)
    p.paragraph_format.space_after  = Pt(12)


def _add_glossary_term(document, text):
    """Bold glossary term."""
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)


def _add_glossary_def(document, text):
    """Indented glossary definition."""
    p = document.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)


def generate_docx(book_text, output_path):

    document = Document()

    _set_margins(document, margin_inches=1.0)

    current_section = None
    in_chapter      = False
    first_element   = True

    lines = book_text.split("\n")

    i = 0

    while i < len(lines):

        line = lines[i]

        # ==========================================
        # SECTION TAG
        # ==========================================

        section_match = re.match(
            r"^\[SECTION:(\w+)\]$",
            line.strip()
        )

        if section_match:

            section_key    = section_match.group(1)
            current_section = section_key
            in_chapter      = False

            if not first_element:
                document.add_page_break()

            first_element = False

            label = SECTION_LABELS.get(section_key)

            if label:
                _add_section_title(document, label)

            i += 1
            continue

        # ==========================================
        # CHAPTER TAG
        # ==========================================

        chapter_match = re.match(
            r"^\[CHAPTER:(\d+)\]$",
            line.strip()
        )

        if chapter_match:

            chapter_num     = chapter_match.group(1)
            in_chapter      = True
            current_section = None

            if not first_element:
                document.add_page_break()

            first_element = False

            _add_chapter_label(document, chapter_num)

            i += 1
            continue

        # ==========================================
        # CHAPTER TITLE TAG
        # ==========================================

        chapter_title_match = re.match(
            r"^\[CHAPTER_TITLE\](.*?)\[/CHAPTER_TITLE\]$",
            line.strip()
        )

        if chapter_title_match:
            _add_chapter_title(
                document,
                chapter_title_match.group(1).strip()
            )
            i += 1
            continue

        # ==========================================
        # GLOSSARY TERM TAG
        # ==========================================

        term_match = re.match(
            r"^\[TERM\](.*?)\[/TERM\]$",
            line.strip()
        )

        if term_match:
            _add_glossary_term(
                document,
                term_match.group(1).strip()
            )
            i += 1
            continue

        # ==========================================
        # GLOSSARY DEFINITION TAG
        # ==========================================

        def_match = re.match(
            r"^\[DEF\](.*?)\[/DEF\]$",
            line.strip()
        )

        if def_match:
            _add_glossary_def(
                document,
                def_match.group(1).strip()
            )
            i += 1
            continue

        # ==========================================
        # BLANK LINE — skip
        # ==========================================

        if not line.strip():
            i += 1
            continue

        # ==========================================
        # CONTENT LINES
        # ==========================================

        text = line.strip()

        # --- Half title / title page ---
        if current_section in ("HALF_TITLE", "TITLE_PAGE"):
            _add_book_title(document, text)
            i += 1
            continue

        # --- Epigraph ---
        if current_section == "EPIGRAPH":
            _add_epigraph(document, text)
            i += 1
            continue

        # --- TOC entries ---
        if current_section == "TOC":
            _add_toc_entry(document, text)
            i += 1
            continue

        # --- Inside a chapter: detect **bold** section headings ---
        if in_chapter:

            bold_match = re.match(r"^\*\*(.*?)\*\*$", text)

            if bold_match:
                _add_section_heading(
                    document,
                    bold_match.group(1)
                )
            else:
                _add_body(document, text)

            i += 1
            continue

        # --- Default: body text ---
        _add_body(document, text)
        i += 1

    document.save(output_path)

    return output_path
